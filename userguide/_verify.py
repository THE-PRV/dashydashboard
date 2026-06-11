from docx import Document
import re
p = r"Y:\dashythedashdashboard\dashythedashdashboard\userguide\Access Review - User Guide.docx"
d = Document(p)
print("paras:", len(d.paragraphs), " tables:", len(d.tables))
print("\n=== HEADINGS ===")
for para in d.paragraphs:
    if para.style.name.startswith("Heading"):
        print(f"  [{para.style.name}] {para.text[:70]}")
print("\n=== dangling 'Section' references ===")
full = "\n".join(p.text for p in d.paragraphs)
for m in re.finditer(r".{0,40}Section(?![ ]?\d)[^\w].{0,30}", full):
    print("  ...", m.group(0).replace("\n"," "))
print("\n=== leftover LaTeX / placeholder check ===")
for needle in ["\\ref","\\label","ref{","Lorem ipsum","Aptos. 20","lorem"]:
    print(f"  {needle!r}: {full.count(needle)}")
print("\n=== callout/table styles ===")
for i,t in enumerate(d.tables):
    rows=len(t.rows); cols=len(t.columns)
    first=t.cell(0,0).text[:40].replace("\n"," ")
    print(f"  table {i}: {rows}x{cols} style={t.style.name if t.style else '-'} :: {first!r}")
