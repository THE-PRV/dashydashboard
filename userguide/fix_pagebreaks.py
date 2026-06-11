# -*- coding: utf-8 -*-
"""
Post-process the generated user guide DOCX to fix page-break problems:
  1. Cap oversized images (>7.0 in tall) preserving aspect ratio.
  2. Keep each figure paragraph with its caption (keepNext + keepLines on image
     para; keepLines on caption para).
  3. Keep headings with their following content (keepNext + keepLines).
  4. Stop tables splitting across pages: cantSplit on every row; tblHeader on the
     first row of each multi-row DATA table.

Does NOT rebuild the document. Preserves cover/TOC/headers/footers/all content.
"""
import sys
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = r"Y:\dashythedashdashboard\dashythedashdashboard\userguide\Access Review - User Guide.docx"
FIXED = r"Y:\dashythedashdashboard\dashythedashdashboard\userguide\Access Review - User Guide (fixed).docx"

EMU_PER_IN = 914400
MAX_H_IN = 7.0


def para_contains_drawing(para):
    """True if this paragraph contains an inline drawing (image)."""
    return para._p.findall(qn('w:r') + '/' + qn('w:drawing')) != [] or \
        bool(para._p.findall('.//' + qn('w:drawing')))


def main():
    doc = Document(SRC)

    # ---- 1. Cap oversized images -------------------------------------------
    rescaled = []
    for i, sh in enumerate(doc.inline_shapes):
        w = sh.width
        h = sh.height
        if h is None or w is None:
            continue
        h_in = h / EMU_PER_IN
        if h_in > MAX_H_IN:
            factor = (MAX_H_IN * EMU_PER_IN) / h
            new_w = int(round(w * factor))
            new_h = int(round(h * factor))
            old_h_in = h_in
            sh.width = new_w
            sh.height = new_h
            rescaled.append((i, old_h_in, new_h / EMU_PER_IN))

    # ---- 2 & 3. keepNext / keepLines on figure + heading paragraphs --------
    paras = doc.paragraphs
    fig_img_count = 0
    fig_caption_count = 0
    heading_count = 0

    for idx, para in enumerate(paras):
        style_name = para.style.name if para.style else ""

        if para_contains_drawing(para):
            # figure image paragraph: keep with next (caption) + keep lines together
            para.paragraph_format.keep_with_next = True
            para.paragraph_format.keep_together = True
            fig_img_count += 1
            # the following caption paragraph: only keepLines
            if idx + 1 < len(paras):
                nxt = paras[idx + 1]
                if nxt.text.startswith("Figure "):
                    nxt.paragraph_format.keep_together = True
                    fig_caption_count += 1

        if style_name.startswith("Heading"):
            para.paragraph_format.keep_with_next = True
            para.paragraph_format.keep_together = True
            heading_count += 1

    # ---- 4. Table row cantSplit + data-table header repeat -----------------
    cantsplit_rows = 0
    tblheader_rows = 0
    for t in doc.tables:
        nrows = len(t.rows)
        ncols = len(t.columns)
        is_data_table = nrows > 1 and ncols > 1  # multi-row data table
        for ri, row in enumerate(t.rows):
            trPr = row._tr.get_or_add_trPr()
            if trPr.find(qn('w:cantSplit')) is None:
                trPr.append(OxmlElement('w:cantSplit'))
                cantsplit_rows += 1
            if is_data_table and ri == 0:
                if trPr.find(qn('w:tblHeader')) is None:
                    trPr.append(OxmlElement('w:tblHeader'))
                    tblheader_rows += 1

    # ---- Save (handle Word lock) -------------------------------------------
    saved_path = SRC
    locked = False
    try:
        doc.save(SRC)
    except PermissionError:
        locked = True
        saved_path = FIXED
        doc.save(FIXED)

    # ---- Report ------------------------------------------------------------
    print("=== APPLIED ===")
    print(f"images total: {len(doc.inline_shapes)}")
    print(f"rescaled images: {len(rescaled)}")
    for i, oh, nh in rescaled:
        print(f"  img {i}: {oh:.2f}in -> {nh:.2f}in")
    print(f"figure image paras (keepNext+keepLines): {fig_img_count}")
    print(f"figure caption paras (keepLines): {fig_caption_count}")
    print(f"heading paras (keepNext+keepLines): {heading_count}")
    print(f"rows given cantSplit: {cantsplit_rows}")
    print(f"data-table header rows given tblHeader: {tblheader_rows}")
    print(f"saved to: {saved_path}")
    print(f"original path writable: {not locked}")
    if locked:
        print("!! ORIGINAL WAS LOCKED (open in Word). Saved to the (fixed) copy instead.")


if __name__ == "__main__":
    main()
