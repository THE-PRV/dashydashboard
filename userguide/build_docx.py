# -*- coding: utf-8 -*-
"""
Build "Access Review - User Guide.docx" from the Broadridge template (_tpl_copy.docx)
+ content from main.tex. Run with python 3.14 + python-docx 1.2.
"""
import os, re, zipfile, shutil
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
TPL  = os.path.join(HERE, "_tpl_copy.docx")
OUT  = os.path.join(HERE, "Access Review - User Guide.docx")
FIGDIR = os.path.join(HERE, "figures")

# ---------------------------------------------------------------------------
# Figure registry: order matches main.tex \fig appearance. width in inches.
# label -> (filename, width_inches)
# ---------------------------------------------------------------------------
NARROW = 3.8
FULL   = 6.3
FIGURES = [
    ("fig:login",       "login-sso.png",            NARROW),
    ("fig:cyclepicker", "assoc-cycle-picker.png",   FULL),
    ("fig:assocdash",   "assoc-dashboard.png",      FULL),
    ("fig:attesttable", "assoc-attest-table.png",   5.0),
    ("fig:remark",      "assoc-remark-modal.png",   NARROW),
    ("fig:submitted",   "assoc-submitted.png",      FULL),
    ("fig:dark",        "dark-mode.png",            FULL),
    ("fig:mgr",         "mgr-overview.png",         FULL),
    ("fig:mgrdetail",   "mgr-member-detail.png",    FULL),
    ("fig:access",      "access-overview.png",      FULL),
    ("fig:grantform",   "access-grant-form.png",    NARROW),
    ("fig:admin",       "admin-dashboard.png",      FULL),
    ("fig:gfh",         "gfh-dashboard.png",        FULL),
    ("fig:delegate",    "gfhdelegate-dashboard.png",FULL),
    ("fig:ifh",         "ifh-dashboard.png",        FULL),
    ("fig:drill",       "admin-drilldown.png",      FULL),
    ("fig:notif",       "admin-notifications.png",  NARROW),
    ("fig:addclient",   "admin-add-client.png",     NARROW),
    ("fig:addtool",     "admin-add-tool.png",       NARROW),
    ("fig:users",       "users-directory.png",      FULL),
    ("fig:useredit",    "users-edit-modal.png",     NARROW),
]
LABEL_NUM = {lab: i + 1 for i, (lab, _, _) in enumerate(FIGURES)}  # label -> figure number

CALLOUT_COLORS = {
    "NOTE":      ("2563EB", "E8EEFC"),
    "TIP":       ("0F9D58", "E7F4EC"),
    "IMPORTANT": ("B45309", "FBF0E2"),
    "ADMIN & OVERSIGHT ROLES": ("000C36", "E7E9EF"),
}

# ---------------------------------------------------------------------------
# Inline LaTeX -> list of (text, style) runs.  style flags: b,i,m
# ---------------------------------------------------------------------------
def resolve_ref(s):
    """Replace \\ref{fig:x}/Section~\\ref{...} etc. We render figure refs as 'Figure N'."""
    # \ref{fig:...} -> the number
    def repl(m):
        lab = m.group(1)
        if lab in LABEL_NUM:
            return str(LABEL_NUM[lab])
        return ""  # section refs -> drop the number (we have no numbering); handled below
    return re.sub(r"\\ref\{([^}]*)\}", repl, s)

def parse_inline(text):
    """Return list of (run_text, set_of_flags). Handles \\ui \\textbf \\tab \\emph \\textit \\texttt and escapes."""
    # Normalize references and section cross-refs first.
    # "Section~\ref{sec:x}" -> "Section" (drop tilde+number) ; "Figure~\ref{fig}" -> "Figure N"
    text = text.replace("~", " ")  # nbsp; we will collapse Section  below
    # Handle \ref
    def ref_repl(m):
        lab = m.group(1)
        if lab in LABEL_NUM:
            return str(LABEL_NUM[lab])
        return ""
    text = re.sub(r"\\ref\{([^}]*)\}", ref_repl, text)
    # collapse "Section " (no number) -> "this guide" style: keep just "Section"
    text = text.replace("Section ", "Section ").replace("Figure ", "Figure ")
    text = re.sub(r"Section \b(?=[A-Z]|$|\.|,|\))", "the relevant section ", text) if False else text
    # Tidy leftover "Section  " double spaces / "Section ."
    text = text.replace("Section .", "the relevant section.").replace("Section )", "the relevant section)")
    text = re.sub(r"\s+", " ", text)
    text = text.replace(" ", " ")

    # Drop structural / no-render commands that take a brace arg we must DISCARD entirely.
    for cmd in ("label", "graphicspath", "hypersetup"):
        text = re.sub(r"\\%s\{[^{}]*\}" % cmd, "", text)
    # Drop color/font-switch commands (keep following text). These take a single brace
    # arg that is a *value*, not content to display.
    text = re.sub(r"\\color\{[^{}]*\}", "", text)
    # Drop bare font-size / spacing switches.
    text = re.sub(r"\\(footnotesize|small|large|Large|Huge|LARGE|normalsize|centering|noindent|clearpage|vfill|thispagestyle\{[^{}]*\})", "", text)
    text = re.sub(r"\\vspace\*?\{[^{}]*\}", "", text)
    text = re.sub(r"\\rule\{[^{}]*\}\{[^{}]*\}", "", text)
    # line breaks -> space
    text = text.replace("\\\\", " ")

    # escapes
    text = text.replace(r"\&", "&").replace(r"\%", "%").replace(r"\#", "#")
    text = text.replace(r"\$", "$").replace(r"\_", "_")
    text = text.replace(r"\,", " ").replace(r"\ ", " ")
    text = text.replace(r"\checkmark", "✓")
    text = text.replace("``", "“").replace("''", "”")
    text = text.replace("---", "—").replace("--", "–")

    # Tokenize commands with a single brace argument.
    runs = []
    i = 0
    n = len(text)
    flag_cmds = {
        r"\ui": "b", r"\textbf": "b", r"\tab": "b", r"\textsf": "",
        r"\emph": "i", r"\textit": "i", r"\texttt": "m", r"\textmd": "",
    }
    def emit(t, flags):
        if t:
            runs.append((t, set(flags)))
    buf = ""
    while i < n:
        if text[i] == "\\":
            # match longest command
            m = re.match(r"\\[a-zA-Z]+", text[i:])
            if m:
                cmd = m.group(0)
                j = i + len(cmd)
                # skip whitespace
                while j < n and text[j] == " ":
                    j += 1
                if j < n and text[j] == "{":
                    # find matching brace
                    depth = 1; k = j + 1; start = k
                    while k < n and depth:
                        if text[k] == "{": depth += 1
                        elif text[k] == "}": depth -= 1
                        k += 1
                    inner = text[start:k-1]
                    flag = flag_cmds.get(cmd, None)
                    if flag is not None:
                        emit(buf, []); buf = ""
                        for sub_t, sub_f in parse_inline_recurse(inner, flag):
                            emit(sub_t, sub_f)
                        i = k
                        continue
                    else:
                        # unknown command with arg: keep inner text plain
                        emit(buf, []); buf = ""
                        for sub_t, sub_f in parse_inline_recurse(inner, ""):
                            emit(sub_t, sub_f)
                        i = k
                        continue
                else:
                    # command without brace arg: drop it
                    i = j
                    continue
            else:
                buf += text[i]; i += 1
                continue
        else:
            buf += text[i]; i += 1
    emit(buf, [])
    # merge adjacent identical-flag runs
    merged = []
    for t, f in runs:
        if merged and merged[-1][1] == f:
            merged[-1][0] = merged[-1][0] + t
        else:
            merged.append([t, f])
    # trim leading/trailing whitespace on the paragraph as a whole
    if merged:
        merged[0][0] = merged[0][0].lstrip()
        merged[-1][0] = merged[-1][0].rstrip()
    merged = [m for m in merged if m[0] != ""]
    return [(t, f) for t, f in merged]

def parse_inline_recurse(inner, extra_flag):
    base = parse_inline(inner)
    out = []
    for t, f in base:
        nf = set(f)
        if extra_flag:
            nf.add(extra_flag)
        out.append((t, nf))
    return out

def add_runs(paragraph, runs):
    for t, flags in runs:
        r = paragraph.add_run(t)
        if "b" in flags:
            r.bold = True
        if "i" in flags:
            r.italic = True
        if "m" in flags:
            r.font.name = "Consolas"
            rpr = r._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
            for a in ("w:ascii", "w:hAnsi", "w:cs"):
                rfonts.set(qn(a), "Consolas")
    return paragraph

# ---------------------------------------------------------------------------
# main.tex parsing
# ---------------------------------------------------------------------------
def load_body():
    with open(os.path.join(HERE, "main.tex"), encoding="utf-8") as f:
        src = f.read()
    # take everything after \clearpage following \tableofcontents up to \end{document}
    start = src.index("\\section{Introduction}")
    end = src.index("\\end{document}")
    return src[start:end]

def strip_comments(s):
    out = []
    for line in s.splitlines():
        # remove unescaped % comments
        res = ""
        esc = False
        for ch in line:
            if ch == "%" and not esc:
                break
            res += ch
            esc = (ch == "\\" and not esc)
        out.append(res)
    return "\n".join(out)

# ---------------------------------------------------------------------------
# Block-level tokenizer
# ---------------------------------------------------------------------------
def tokenize(body):
    """Yield block tokens as dicts."""
    body = strip_comments(body)
    tokens = []
    i = 0
    n = len(body)
    env_re = re.compile(r"\\begin\{(itemize|enumerate|description|table|figure|center|notebox|tipbox|importantbox|adminbox|tabularx|tabular)\}")
    while i < n:
        # find next control of interest
        m = re.search(r"\\(section|subsection|subsubsection)\{", body[i:])
        e = re.search(r"\\begin\{(\w+)\}", body[i:])
        f = re.search(r"\\fig\{", body[i:])
        candidates = []
        if m: candidates.append((i + m.start(), "head", m))
        if e: candidates.append((i + e.start(), "env", e))
        if f: candidates.append((i + f.start(), "fig", f))
        if not candidates:
            # rest is body text
            txt = body[i:].strip()
            if txt:
                for para in split_paras(txt):
                    tokens.append({"t": "para", "text": para})
            break
        candidates.sort()
        pos, kind, mm = candidates[0]
        # text before
        pre = body[i:pos].strip()
        if pre:
            for para in split_paras(pre):
                tokens.append({"t": "para", "text": para})
        if kind == "head":
            name = mm.group(1)
            argstart = pos + len(mm.group(0))
            arg, after = read_braced(body, argstart - 1)
            tokens.append({"t": name, "text": arg})
            i = after
        elif kind == "fig":
            # \fig{file}{width}{caption}{label}
            file_, p = read_braced(body, pos + len("\\fig"))
            width_, p = read_braced(body, p)
            cap_, p = read_braced(body, p)
            lab_, p = read_braced(body, p)
            tokens.append({"t": "fig", "file": file_.strip(), "caption": cap_, "label": lab_.strip()})
            i = p
        else:  # env
            envname = mm.group(1)
            envstart = pos
            content, after = read_env(body, pos, envname)
            tokens.append({"t": "env", "env": envname, "content": content})
            i = after
    return tokens

def split_paras(text):
    parts = re.split(r"\n\s*\n", text)
    return [re.sub(r"\s+", " ", p).strip() for p in parts if p.strip()]

def read_braced(s, idx):
    """idx points at or before '{'. Returns (inner, index_after_closing)."""
    while s[idx] != "{":
        idx += 1
    depth = 1
    k = idx + 1
    start = k
    while k < len(s) and depth:
        if s[k] == "{": depth += 1
        elif s[k] == "}": depth -= 1
        k += 1
    return s[start:k-1], k

def read_env(s, pos, envname):
    begin = "\\begin{%s}" % envname
    end = "\\end{%s}" % envname
    bstart = s.index(begin, pos) + len(begin)
    # handle nested same-env (none expected except itemize within itemize - handled by caller)
    depth = 1
    k = bstart
    while depth:
        nb = s.find(begin, k)
        ne = s.find(end, k)
        if ne == -1:
            break
        if nb != -1 and nb < ne:
            depth += 1; k = nb + len(begin)
        else:
            depth -= 1
            if depth == 0:
                return s[bstart:ne], ne + len(end)
            k = ne + len(end)
    return s[bstart:], len(s)

# ---------------------------------------------------------------------------
# list / description item splitting
# ---------------------------------------------------------------------------
def split_items(content):
    """Split an itemize/enumerate body into items by top-level \\item, allowing nested envs."""
    items = []
    # find \item positions at top level (not inside nested begin/end)
    idxs = []
    depth = 0
    i = 0
    while i < len(content):
        if content.startswith("\\begin{", i):
            depth += 1; i += 7; continue
        if content.startswith("\\end{", i):
            depth -= 1; i += 5; continue
        if depth == 0 and content.startswith("\\item", i):
            idxs.append(i)
            i += 5; continue
        i += 1
    for n, st in enumerate(idxs):
        en = idxs[n+1] if n+1 < len(idxs) else len(content)
        items.append(content[st+5:en])
    return items

def parse_description_item(item):
    """Return (term_raw, body_raw). item begins after \\item, optional [term]."""
    item = item.lstrip()
    if item.startswith("["):
        depth = 1; k = 1
        while k < len(item) and depth:
            if item[k] == "[": depth += 1
            elif item[k] == "]": depth -= 1
            k += 1
        term = item[1:k-1]
        body = item[k:]
        return term, body
    return None, item

# ---------------------------------------------------------------------------
# Word emission helpers
# ---------------------------------------------------------------------------
def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def set_cell_borders(cell, color="B6C0CC", sz="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), sz)
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)

def set_table_width_pct(table, pct=100):
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"), str(pct * 50))

CONTENT_WIDTH = Inches(6.5)

# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------
def remove_showcase(doc):
    body = doc.element.body
    children = list(body)
    # sectPr is the last child
    sectPr = body.find(qn("w:sectPr"))
    # Showcase starts right after the TOC field end. The TOC field 'end' fldChar
    # lives in the paragraph indexed where the heading placeholders begin.
    # We identified: keep indices 0..16 (cover, blanks, TOC heading, TOC field incl its end para 16).
    # Delete from 17 up to but not including sectPr.
    # Robust approach: find the TOC field end fldChar, delete everything after its paragraph
    # until sectPr.
    # Find paragraph containing fldChar end that closes TOC (last 'end' belonging to TOC instr).
    # Simpler: locate first placeholder paragraph that begins the showcase by text marker.
    start_idx = None
    for idx, el in enumerate(children):
        if el.tag == qn("w:p"):
            txt = "".join(t.text or "" for t in el.iter(qn("w:t")))
            if "Heading 1. Aptos. 20 Points" in txt:
                start_idx = idx
                break
    assert start_idx is not None, "could not find showcase start"
    # But para 16 (empty Heading1 with TOC end) and para 17 (empty) precede it; those
    # are leftover blanks/TOC-end. Keep the TOC-end paragraph (16) but drop the trailing
    # empty paragraph (17) and everything from start_idx onward up to sectPr.
    # Remove from index 17 .. sectPr (exclusive). 17 is the empty para right after TOC end.
    # Determine index of sectPr:
    sect_idx = children.index(sectPr)
    # We delete children[17 : sect_idx]
    for el in children[17:sect_idx]:
        body.remove(el)
    return doc

def add_para(doc, runs=None, style=None, align=None, size=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if runs:
        add_runs(p, runs)
    if size:
        for r in p.runs:
            r.font.size = Pt(size)
    return p

def emit_callout(doc, label, body_text):
    color, fill = CALLOUT_COLORS[label]
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width_pct(tbl, 100)
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_borders(cell, color="C9D2DD", sz="6")
    # label paragraph
    cell.paragraphs[0].text = ""
    lp = cell.paragraphs[0]
    lr = lp.add_run(label)
    lr.bold = True
    lr.font.color.rgb = RGBColor.from_string(color)
    lr.font.size = Pt(9)
    # body paragraph
    bp = cell.add_paragraph()
    add_runs(bp, parse_inline(body_text))
    return tbl

def emit_figure(doc, file_, caption, label):
    num = LABEL_NUM[label]
    fname, width = next((fn, w) for (lb, fn, w) in FIGURES if lb == label)
    path = os.path.join(FIGDIR, fname)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    # caption
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run("Figure %d — " % num)
    cr.bold = True
    cr.font.size = Pt(9)
    cap_runs = parse_inline(caption)
    for t, flags in cap_runs:
        r = cp.add_run(t)
        r.italic = True
        r.font.size = Pt(9)
        if "b" in flags:
            r.bold = True

def emit_table_from_tabularx(doc, content, header_is_first=True):
    """Parse a simple tabularx/tabular body (rows separated by \\\\, cols by &)."""
    # strip \toprule \midrule \bottomrule and \renewcommand etc already handled by caller
    content = re.sub(r"\\(toprule|midrule|bottomrule)", "", content)
    rows = []
    for raw in re.split(r"\\\\", content):
        raw = raw.strip()
        if not raw:
            continue
        cells = split_top_level(raw, "&")
        rows.append([c.strip() for c in cells])
    ncol = max(len(r) for r in rows)
    tbl = doc.add_table(rows=0, cols=ncol)
    try:
        tbl.style = "Grid Table Light"
    except KeyError:
        tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width_pct(tbl, 100)
    for ri, row in enumerate(rows):
        wrow = tbl.add_row()
        for ci in range(ncol):
            cell = wrow.cells[ci]
            cell.paragraphs[0].text = ""
            txt = row[ci] if ci < len(row) else ""
            runs = parse_inline(txt)
            # bold header row
            if ri == 0 and header_is_first:
                runs = [(t, f | {"b"}) for t, f in runs] if runs else [("", {"b"})]
            add_runs(cell.paragraphs[0], runs)
    return tbl

def split_top_level(s, sep):
    out = []
    depth = 0
    cur = ""
    i = 0
    while i < len(s):
        ch = s[i]
        if ch == "{":
            depth += 1; cur += ch
        elif ch == "}":
            depth -= 1; cur += ch
        elif ch == "\\" and i+1 < len(s):
            cur += s[i:i+2]; i += 2; continue
        elif ch == sep and depth == 0:
            out.append(cur); cur = ""
        else:
            cur += ch
        i += 1
    out.append(cur)
    return out

def emit_env(doc, env, content):
    if env in ("notebox", "tipbox", "importantbox", "adminbox"):
        label = {"notebox": "NOTE", "tipbox": "TIP",
                 "importantbox": "IMPORTANT", "adminbox": "ADMIN & OVERSIGHT ROLES"}[env]
        text = re.sub(r"\s+", " ", content).strip()
        emit_callout(doc, label, text)
    elif env == "itemize":
        emit_list(doc, content, "List Bullet", "List Bullet 2", numbered=False)
    elif env == "enumerate":
        emit_list(doc, content, "List Number", "List Number 2", numbered=True)
    elif env == "description":
        emit_description(doc, content)
    elif env in ("table", "figure", "center"):
        # could contain tabularx or a centered \texttt etc.
        if "tabularx" in content or "tabular" in content:
            inner = extract_inner_env(content)
            emit_table_from_tabularx(doc, inner)
        else:
            text = re.sub(r"\s+", " ", content).strip()
            if text:
                p = add_para(doc, parse_inline(text), style="Normal",
                             align=WD_ALIGN_PARAGRAPH.CENTER)
    elif env in ("tabularx", "tabular"):
        emit_table_from_tabularx(doc, strip_tabular_spec(content))

def extract_inner_env(content):
    m = re.search(r"\\begin\{tabularx\}", content)
    if m:
        inner, _ = read_env(content, 0, "tabularx")
        return strip_tabular_spec(inner)
    inner, _ = read_env(content, 0, "tabular")
    return strip_tabular_spec(inner)

def strip_tabular_spec(inner):
    """tabularx body begins with {\\linewidth}{colspec}. Remove the two leading brace groups."""
    inner = inner.strip()
    # remove {\linewidth}
    if inner.startswith("{"):
        _, p = read_braced(inner, 0)
        inner = inner[p:].lstrip()
        if inner.startswith("{"):
            _, p2 = read_braced(inner, 0)
            inner = inner[p2:]
    return inner

def emit_list(doc, content, style1, style2, numbered):
    items = split_items(content)
    for item in items:
        # check for nested itemize
        nested = None
        nm = re.search(r"\\begin\{itemize\}", item)
        if nm:
            head = item[:nm.start()]
            nested_content, _ = read_env(item, nm.start(), "itemize")
            item_text = re.sub(r"\s+", " ", head).strip()
            if item_text:
                add_para(doc, parse_inline(item_text), style=style1)
            for sub in split_items(nested_content):
                st = re.sub(r"\s+", " ", sub).strip()
                if st:
                    add_para(doc, parse_inline(st), style=style2)
        else:
            st = re.sub(r"\s+", " ", item).strip()
            if st:
                add_para(doc, parse_inline(st), style=style1)

def emit_description(doc, content):
    items = split_items(content)
    for item in items:
        term, body = parse_description_item(item)
        body = re.sub(r"\s+", " ", body).strip()
        p = doc.add_paragraph(style="Normal")
        if term is not None:
            term_runs = parse_inline(term)
            # term in bold
            for t, f in term_runs:
                r = p.add_run(t)
                r.bold = True
            p.add_run(" — ")
        add_runs(p, parse_inline(body))

# ---------------------------------------------------------------------------
def set_cover_title(doc):
    doc.core_properties.title = "Access Review - User Guide"
    doc.core_properties.subject = "Client Application Access Attestation - BPO Operations"
    doc.core_properties.author = "Broadridge BPO"
    doc.core_properties.category = "User Guide"

    body = doc.element.body
    sdt = body[0]
    # Replace the title placeholder runs: the title spans many w:t runs forming
    # "Title, title case maximum two-line title". We collapse them into one.
    # Strategy: find the contiguous run group whose concatenation contains "Title, title case".
    all_t = list(sdt.iter(qn("w:t")))
    full = "".join(t.text or "" for t in all_t)
    # locate title and subtitle text nodes
    # Title nodes: from first until we reach the 'Subtitle' node
    title_set = False
    sub_set = False
    for idx, t in enumerate(all_t):
        txt = t.text or ""
        if not title_set and txt.startswith("Title, "):
            # this is the first title node; set it and blank the rest of the title fragment
            t.text = "Access Review — User Guide"
            title_set = True
            continue
        if title_set and not sub_set and txt and not txt.startswith("Subtitle"):
            # subsequent title fragment nodes -> blank, until subtitle reached
            # but stop blanking once we hit 'Subtitle'
            t.text = ""
            continue
        if txt.startswith("Subtitle"):
            t.text = "Client Application Access Attestation — BPO Operations"
            sub_set = True
            continue
        if sub_set and txt == "placeholder one line":
            t.text = ""
            continue
    return doc

# ---------------------------------------------------------------------------
def inject_update_fields(path):
    tmp = path + ".tmp"
    with zipfile.ZipFile(path, "r") as zin:
        names = zin.namelist()
        data = {n: zin.read(n) for n in names}
    settings = data["word/settings.xml"].decode("utf-8")
    if "w:updateFields" not in settings:
        m = re.search(r"<w:settings[^>]*>", settings)
        insert = '<w:updateFields w:val="true"/>'
        settings = settings[:m.end()] + insert + settings[m.end():]
        data["word/settings.xml"] = settings.encode("utf-8")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for n in names:
            zout.writestr(n, data[n])
    shutil.move(tmp, path)

# ---------------------------------------------------------------------------
def main():
    doc = Document(TPL)
    remove_showcase(doc)
    set_cover_title(doc)

    body = load_body()
    tokens = tokenize(body)

    style_map = {"section": "Heading 1", "subsection": "Heading 2", "subsubsection": "Heading 3"}

    for tok in tokens:
        t = tok["t"]
        if t in style_map:
            add_para(doc, parse_inline(tok["text"]), style=style_map[t])
        elif t == "para":
            txt = tok["text"]
            # drop standalone end-of-guide rule artefacts handled as plain text
            add_para(doc, parse_inline(txt), style="Normal")
        elif t == "fig":
            emit_figure(doc, tok["file"], tok["caption"], tok["label"])
        elif t == "env":
            emit_env(doc, tok["env"], tok["content"])

    doc.save(OUT)
    inject_update_fields(OUT)
    print("saved", OUT)

if __name__ == "__main__":
    main()
